"""Template-fill (mail-merge) helper.

Turn ONE master template (a .docx letter/form or a .xlsx form) plus a data table
(one row per output) into MANY filled copies — preserving the master's layout and
branding exactly, swapping only tokens.

Token syntax: ``{{TokenName}}`` (whitespace inside the braces is ignored).

The flow the SKILL drives, and the functions that serve each step:

    1. ANALYSE   read_content(path)            -> the master's text, so you can
                                                  spot the parts that vary.
    2. TOKENISE  tokenise(src, dst, mapping)   -> write a tokenised copy: each
                                                  variable phrase becomes {{Token}}.
                 tokens_in(path)               -> list the tokens now in the template.
    3. DATA      load_rows(path)               -> (headers, [row-dict, ...]) from
                                                  .xlsx / .csv.
    4. GENERATE  generate(tmpl, rows, ...)     -> one filled file per row + a report.

Design rules (from SKILL.md ## Principles):
  * Deterministic — same inputs, same outputs. No network, no model calls here.
  * Never invent — a token with no data for a row is written as a VISIBLE flag
    («MISSING: Token»), never a silent blank, and is listed in the report.
  * Preserve the master — replacement edits only the runs/cells a token occupies;
    formatting elsewhere is untouched.

Usage as a library:

    from fill_template import read_content, tokenise, tokens_in, load_rows, generate

    # 1. analyse
    print(read_content("ConfirmationLetter_master.docx"))

    # 2. tokenise a copy (mapping confirmed with the user)
    tokenise("ConfirmationLetter_master.docx", "ConfirmationLetter_tokenised.docx",
             [{"find": "Ms Jordan Lee", "token": "RecipientName"},
              {"find": "$1,000,000",  "token": "Amount"}])

    # 3. load the data table
    headers, rows = load_rows("recipients.xlsx")

    # 4. generate one file per row
    report = generate(
        "ConfirmationLetter_tokenised.docx", rows,
        token_to_column={"RecipientName": "Name"},
        outdir="out",
        name_pattern="ConfirmationLetter_{Name}",
    )
"""

from __future__ import annotations

import csv
import datetime
import re
from collections import Counter
from pathlib import Path

import openpyxl
from docx import Document

# ----------------------------------------------------------------------------
# Token plumbing
# ----------------------------------------------------------------------------

TOKEN_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _token(name: str) -> str:
    return "{{" + name + "}}"


def _missing(name: str) -> str:
    # Visible, greppable, never a silent blank.
    return f"«MISSING: {name}»"


def _try_parse_date(s: str) -> datetime.date | None:
    """Try to parse a string as an ISO date (YYYY-MM-DD or YYYY/MM/DD).

    CSV data comes in as strings, not datetime objects — this lets _fmt_value
    still render them as DD MMM YYYY. Returns None if the string isn't a date
    (so non-date strings pass through untouched).
    """
    s = s.strip()
    if not s:
        return None
    # Accept YYYY-MM-DD or YYYY/MM/DD (the common ISO / spreadsheet export forms).
    for sep in ("-", "/"):
        parts = s.split(sep)
        if len(parts) == 3:
            try:
                y, m, d = int(parts[0]), int(parts[1]), int(parts[2])
                # Sanity-check month/day so "2026-13-01" doesn't silently parse.
                if 1 <= m <= 12 and 1 <= d <= 31:
                    return datetime.date(y, m, d)
            except (ValueError, TypeError):
                pass
    return None


def _fmt_value(v) -> str:
    """House-style scalar formatting for a data value.

    Dates -> DD MMM YYYY (datetime objects from .xlsx, and ISO date strings from
    .csv like '2026-07-01'); whole-number floats lose the trailing .0; everything
    else is str(). Currency/precision formatting is the caller's job (format the
    column in the data file) — this only handles the obvious cases.
    """
    if v is None:
        return ""
    if isinstance(v, (datetime.datetime, datetime.date)):
        return v.strftime("%d %b %Y")
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    # CSV dates arrive as strings — try to parse common ISO date formats so they
    # render as DD MMM YYYY, matching .xlsx datetime behaviour. Non-date strings
    # (names, references, amounts) fall through to str() untouched.
    if isinstance(v, str):
        parsed = _try_parse_date(v)
        if parsed is not None:
            return parsed.strftime("%d %b %Y")
    return str(v)


# ----------------------------------------------------------------------------
# Run-aware literal replacement (the .docx hard part)
# ----------------------------------------------------------------------------

def _find_matches(text: str, finds: list[str]) -> list[tuple[int, int, str]]:
    """Left-to-right, non-overlapping matches of any literal in *finds*.

    Returns (start, end, matched_find). At each position the LONGEST matching
    literal wins, so "{{A}}" is preferred over a stray "{{".
    """
    matches: list[tuple[int, int, str]] = []
    i, n = 0, len(text)
    while i < n:
        hit = None
        for f in finds:
            if f and text.startswith(f, i) and (hit is None or len(f) > len(hit)):
                hit = f
        if hit:
            matches.append((i, i + len(hit), hit))
            i += len(hit)
        else:
            i += 1
    return matches


def _replace_in_paragraph(paragraph, repl_for: dict) -> int:
    """Replace literals in one paragraph, preserving run formatting.

    *repl_for* maps a literal -> its replacement string. Matches are located in
    the paragraph's concatenated run text, then applied right-to-left so earlier
    offsets stay valid. The replacement text lands in the run where the match
    begins (inheriting that run's formatting); any other runs the match spans are
    trimmed. Runs outside every match are left completely untouched.
    """
    runs = paragraph.runs
    if not runs:
        return 0
    text = "".join(r.text for r in runs)
    matches = _find_matches(text, list(repl_for.keys()))
    if not matches:
        return 0
    for start, end, found in reversed(matches):
        repl = repl_for[found]
        # Current run spans (recomputed each time; only runs at/after `start`
        # can have changed, so earlier offsets remain valid).
        pos = 0
        first = True
        for r in runs:
            r_start, r_end = pos, pos + len(r.text)
            pos = r_end
            if r_end <= start or r_start >= end:
                continue  # this run does not overlap the match
            local_start = max(r_start, start) - r_start
            local_end = min(r_end, end) - r_start
            before, after = r.text[:local_start], r.text[local_end:]
            if first:
                r.text = before + repl + after
                first = False
            else:
                r.text = before + after
    return len(matches)


def _iter_paragraphs(doc_or_container):
    """Yield every paragraph in a document/cell: body, tables (recursively),
    and — for a Document — each section's header and footer."""
    for p in doc_or_container.paragraphs:
        yield p
    for table in doc_or_container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)
    # Headers/footers only exist on a Document, not on a _Cell.
    sections = getattr(doc_or_container, "sections", None)
    if sections is not None:
        for section in sections:
            for hf in (section.header, section.first_page_header,
                       section.even_page_header, section.footer,
                       section.first_page_footer, section.even_page_footer):
                for p in hf.paragraphs:
                    yield p
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from _iter_paragraphs(cell)


# ----------------------------------------------------------------------------
# 1. ANALYSE
# ----------------------------------------------------------------------------

def read_content(path: str) -> str:
    """Return the master's readable text so the variable parts can be spotted.

    .docx -> paragraph and table text (and header/footer text), one block per line.
    Repeated paragraph text is shown once with a ``(×N)`` count so the user sees
    every location the ``tokenise`` hits will touch (e.g. a name in both the body
    and the header). .xlsx -> each non-empty cell as ``Sheet!A1: value`` (cells are
    where tokens go).
    """
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        doc = Document(path)
        counts: Counter = Counter()
        for p in _iter_paragraphs(doc):
            t = p.text.strip()
            if t:
                counts[t] += 1
        lines = []
        for text, n in counts.items():
            lines.append(text if n == 1 else f"{text}  (×{n})")
        return "\n".join(lines)
    if ext == ".xlsx":
        wb = openpyxl.load_workbook(path, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value not in (None, ""):
                        lines.append(f"{ws.title}!{cell.coordinate}: {cell.value}")
        return "\n".join(lines)
    raise ValueError(f"Unsupported template type: {ext} (use .docx or .xlsx)")


# ----------------------------------------------------------------------------
# 2. TOKENISE
# ----------------------------------------------------------------------------

def tokenise(src: str, dst: str, mapping: list[dict]) -> dict:
    """Write a tokenised copy of *src* to *dst*.

    *mapping* is a list of ``{"find": <exact phrase in the master>, "token":
    <TokenName>}``. Longer phrases are applied first so a phrase that contains a
    shorter one is not pre-empted. Returns a report: per-token hit counts and any
    phrase that was not found (so it can be corrected before generating).
    """
    repl_for = {}
    # Longest find first => correct precedence inside a single paragraph/cell.
    for m in sorted(mapping, key=lambda d: len(d["find"]), reverse=True):
        repl_for[m["find"]] = _token(m["token"])
    counts = {m["token"]: 0 for m in mapping}
    find_to_token = {m["find"]: m["token"] for m in mapping}

    ext = Path(src).suffix.lower()
    if ext == ".docx":
        doc = Document(src)
        for p in _iter_paragraphs(doc):
            text = "".join(r.text for r in p.runs)
            for start, end, found in _find_matches(text, list(repl_for.keys())):
                counts[find_to_token[found]] += 1
            _replace_in_paragraph(p, repl_for)
        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        doc.save(dst)
    elif ext == ".xlsx":
        wb = openpyxl.load_workbook(src)
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        new = cell.value
                        for find in repl_for:
                            if find in new:
                                counts[find_to_token[find]] += new.count(find)
                                new = new.replace(find, repl_for[find])
                        if new != cell.value:
                            cell.value = new
        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        wb.save(dst)
    else:
        raise ValueError(f"Unsupported template type: {ext} (use .docx or .xlsx)")

    not_found = [m["token"] for m in mapping if counts[m["token"]] == 0]
    return {"dst": dst, "hits": counts, "not_found": not_found}


def tokens_in(path: str) -> list[str]:
    """List the distinct tokens present in a (tokenised) template, in first-seen
    order — used to validate the token->column map before generating."""
    text = read_content(path)
    seen, out = set(), []
    for name in TOKEN_RE.findall(text):
        name = name.strip()
        if name not in seen:
            seen.add(name)
            out.append(name)
    return out


# ----------------------------------------------------------------------------
# 3. DATA
# ----------------------------------------------------------------------------

def load_rows(path: str) -> tuple[list[str], list[dict]]:
    """Load the data table. First row = headers; each later row = one output.

    .xlsx -> first worksheet, values only (formulas evaluated by the last app to
    save the file). .csv -> DictReader. Returns (headers, [row-dict, ...]).
    """
    ext = Path(path).suffix.lower()
    if ext == ".xlsx":
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header_row = next(rows_iter)
        except StopIteration:
            return [], []
        headers = [str(h).strip() if h is not None else "" for h in header_row]
        out = []
        for raw in rows_iter:
            if all(c is None or str(c).strip() == "" for c in raw):
                continue  # skip wholly blank rows
            out.append({headers[i]: raw[i] if i < len(raw) else None
                        for i in range(len(headers))})
        return headers, out
    if ext == ".csv":
        with open(path, newline="", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            headers = [h.strip() for h in (reader.fieldnames or [])]
            out = [dict(r) for r in reader
                   if any((v or "").strip() for v in r.values())]
        return headers, out
    raise ValueError(f"Unsupported data type: {ext} (use .xlsx or .csv)")


# ----------------------------------------------------------------------------
# 4. GENERATE
# ----------------------------------------------------------------------------

_ILLEGAL = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def _safe_name(stem: str) -> str:
    stem = _ILLEGAL.sub("_", stem).strip().strip(".")
    stem = re.sub(r"\s+", " ", stem)
    return stem or "output"


def _resolve_map(template_tokens: list[str], token_to_column: dict | None,
                 headers: list[str]) -> dict:
    """Default mapping: token name == column header (case-insensitive). An
    explicit *token_to_column* overrides per token."""
    token_to_column = dict(token_to_column or {})
    lower = {h.lower(): h for h in headers}
    for tok in template_tokens:
        if tok not in token_to_column and tok.lower() in lower:
            token_to_column[tok] = lower[tok.lower()]
    return token_to_column


def _fill_one(template_path: str, values: dict, out_path: str) -> list[str]:
    """Write one filled copy. *values* maps token -> already-stringified value
    (missing tokens absent => flagged). Returns the list of tokens left missing."""
    ext = Path(template_path).suffix.lower()
    tmpl_tokens = tokens_in(template_path)
    repl_for, missing = {}, []
    for tok in tmpl_tokens:
        if tok in values and values[tok] != "":
            repl_for[_token(tok)] = values[tok]
        else:
            repl_for[_token(tok)] = _missing(tok)
            missing.append(tok)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    if ext == ".docx":
        doc = Document(template_path)
        for p in _iter_paragraphs(doc):
            _replace_in_paragraph(p, repl_for)
        doc.save(out_path)
    elif ext == ".xlsx":
        wb = openpyxl.load_workbook(template_path)
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and "{{" in cell.value:
                        new = cell.value
                        for t, r in repl_for.items():
                            new = new.replace(t, r)
                        cell.value = new
        wb.save(out_path)
    else:
        raise ValueError(f"Unsupported template type: {ext} (use .docx or .xlsx)")
    return missing


def generate(template_path: str, rows: list[dict], token_to_column: dict | None = None,
             outdir: str = "out", name_pattern: str | None = None) -> dict:
    """Generate one filled file per row.

    template_path  the TOKENISED master (.docx or .xlsx).
    rows           list of row-dicts from load_rows().
    token_to_column  optional {token: column}. Tokens not given here default to
                   the column whose header matches the token name (case-insensitive).
    outdir         output folder (created if absent).
    name_pattern   output stem, with {Column} placeholders filled from the row,
                   e.g. "ConfirmationLetter_{Name}". Defaults to
                   "<template-stem>_<n>". The template's extension is kept.

    Returns a report:
      {"written":[{"file","row","missing":[...]}], "skipped":[...],
       "unmapped_tokens":[...], "rows":N, "outdir":...}
    Never invents data: an unmapped token, or a mapped column that is blank for a
    row, is written as «MISSING: Token» and recorded.
    """
    ext = Path(template_path).suffix.lower()
    tmpl_tokens = tokens_in(template_path)
    headers = list(rows[0].keys()) if rows else []
    mapping = _resolve_map(tmpl_tokens, token_to_column, headers)
    unmapped = [t for t in tmpl_tokens if t not in mapping]

    outdir_p = Path(outdir)
    outdir_p.mkdir(parents=True, exist_ok=True)
    stem_default = Path(template_path).stem.replace("_tokenised", "")

    written, skipped, used_names = [], [], set()
    for n, row in enumerate(rows, start=1):
        values = {tok: _fmt_value(row.get(mapping[tok]))
                  for tok in tmpl_tokens if tok in mapping}
        # filename
        if name_pattern:
            try:
                stem = name_pattern.format(**{k: _fmt_value(v) for k, v in row.items()})
            except KeyError as e:
                skipped.append({"row": n, "reason": f"name_pattern column {e} not in data"})
                continue
        else:
            stem = f"{stem_default}_{n}"
        stem = _safe_name(stem)
        candidate, k = stem, 2
        while candidate in used_names:  # avoid clobbering same-named rows
            candidate = f"{stem} ({k})"
            k += 1
        used_names.add(candidate)

        out_path = outdir_p / f"{candidate}{ext}"
        missing = _fill_one(template_path, values, str(out_path))
        written.append({"file": str(out_path), "row": n, "missing": missing})

    return {
        "written": written,
        "skipped": skipped,
        "unmapped_tokens": unmapped,
        "rows": len(rows),
        "outdir": str(outdir_p),
    }
