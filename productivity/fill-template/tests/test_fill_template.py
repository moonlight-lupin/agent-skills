#!/usr/bin/env python3
"""Unit tests for fill_template.py — the mail-merge engine.

Run:  python -m pytest productivity/fill-template/tests/test_fill_template.py -v
  or:  python productivity/fill-template/tests/test_fill_template.py

These tests build small synthetic .docx/.xlsx fixtures in a temp dir, so they
don't need any external template files. No network, no credentials.
"""

import sys
import os
import csv
import datetime
import tempfile
import unittest
from pathlib import Path

# Import the module under test (scripts/ is one level up from tests/)
SCRIPTS_DIR = Path(__file__).resolve().parent.parent / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

import fill_template as ft  # noqa: E402
from docx import Document  # noqa: E402
import openpyxl  # noqa: E402


# ---------------------------------------------------------------------------
# Helpers: build small synthetic fixtures
# ---------------------------------------------------------------------------

def _make_docx(path, paragraphs, *, header=None, footer=None):
    """Create a .docx with the given paragraph strings. Optional header/footer text."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if header or footer:
        section = doc.sections[0]
        if header:
            section.header.add_paragraph(header)
        if footer:
            section.footer.add_paragraph(footer)
    doc.save(str(path))


def _make_docx_with_bold(path, before, bold_text, after):
    """Create a .docx where bold_text is in a bold run: [before][bold_text][after]."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(before)
    r = p.add_run(bold_text)
    r.bold = True
    p.add_run(after)
    doc.save(str(path))


def _make_xlsx(path, cells, *, formula_cell=None):
    """Create an .xlsx where cells is {coord: value}. Optional formula in a coord."""
    wb = openpyxl.Workbook()
    ws = wb.active
    for coord, val in cells.items():
        ws[coord] = val
    if formula_cell:
        coord, formula = formula_cell
        ws[coord] = formula
    wb.save(str(path))


def _make_data_csv(path, headers, rows):
    """Write a CSV data file."""
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(str(path), "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        w.writerow(headers)
        for row in rows:
            w.writerow(row)


def _make_data_xlsx(path, headers, rows):
    """Write an XLSX data file. Row values can include datetime.date for real date cells."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(str(path))


# ---------------------------------------------------------------------------
# _fmt_value — scalar formatting
# ---------------------------------------------------------------------------

class TestFmtValue(unittest.TestCase):
    """_fmt_value: dates, floats, zero, None, strings."""

    def test_none(self):
        self.assertEqual(ft._fmt_value(None), "")

    def test_int(self):
        self.assertEqual(ft._fmt_value(42), "42")

    def test_float_whole(self):
        """Whole-number floats lose the trailing .0."""
        self.assertEqual(ft._fmt_value(1000000.0), "1000000")

    def test_float_fractional(self):
        self.assertEqual(ft._fmt_value(3.14), "3.14")

    def test_zero(self):
        """Zero renders as '0', not '' (it's a real value, not missing)."""
        self.assertEqual(ft._fmt_value(0), "0")
        self.assertEqual(ft._fmt_value(0.0), "0")

    def test_datetime_object(self):
        d = datetime.datetime(2026, 7, 1, 0, 0)
        self.assertEqual(ft._fmt_value(d), "01 Jul 2026")

    def test_date_object(self):
        d = datetime.date(2026, 7, 1)
        self.assertEqual(ft._fmt_value(d), "01 Jul 2026")

    def test_csv_date_string_iso_dash(self):
        """CSV dates arrive as strings — should parse to DD MMM YYYY."""
        self.assertEqual(ft._fmt_value("2026-07-01"), "01 Jul 2026")

    def test_csv_date_string_iso_slash(self):
        self.assertEqual(ft._fmt_value("2026/07/01"), "01 Jul 2026")

    def test_non_date_string_passthrough(self):
        """Strings that aren't dates pass through untouched."""
        self.assertEqual(ft._fmt_value("REF-0001"), "REF-0001")
        self.assertEqual(ft._fmt_value("Ms Jordan Lee"), "Ms Jordan Lee")
        self.assertEqual(ft._fmt_value("1000000"), "1000000")

    def test_already_formatted_date_passthrough(self):
        """'01 Jul 2026' is not an ISO date string — passes through as-is."""
        self.assertEqual(ft._fmt_value("01 Jul 2026"), "01 Jul 2026")

    def test_invalid_month_passthrough(self):
        """Invalid month (13) doesn't parse — passes through as-is."""
        self.assertEqual(ft._fmt_value("2026-13-01"), "2026-13-01")

    def test_invalid_day_passthrough(self):
        """Invalid day (32) doesn't parse — passes through as-is."""
        self.assertEqual(ft._fmt_value("2026-07-32"), "2026-07-32")

    def test_empty_string(self):
        self.assertEqual(ft._fmt_value(""), "")

    def test_reference_with_dashes_not_date(self):
        """REF-0001 has dashes but isn't a date — must not be parsed."""
        self.assertEqual(ft._fmt_value("REF-0001"), "REF-0001")


# ---------------------------------------------------------------------------
# _try_parse_date — the ISO date string parser
# ---------------------------------------------------------------------------

class TestTryParseDate(unittest.TestCase):
    """_try_parse_date: ISO date string detection."""

    def test_valid_iso_dash(self):
        self.assertEqual(ft._try_parse_date("2026-07-01"), datetime.date(2026, 7, 1))

    def test_valid_iso_slash(self):
        self.assertEqual(ft._try_parse_date("2026/07/01"), datetime.date(2026, 7, 1))

    def test_invalid_month(self):
        self.assertIsNone(ft._try_parse_date("2026-13-01"))

    def test_invalid_day(self):
        self.assertIsNone(ft._try_parse_date("2026-07-32"))

    def test_not_a_date(self):
        self.assertIsNone(ft._try_parse_date("REF-0001"))
        self.assertIsNone(ft._try_parse_date("hello"))
        self.assertIsNone(ft._try_parse_date("1000000"))

    def test_empty(self):
        self.assertIsNone(ft._try_parse_date(""))

    def test_two_parts(self):
        """Only two parts (YYYY-MM) is not a full date."""
        self.assertIsNone(ft._try_parse_date("2026-07"))

    def test_four_parts(self):
        """Four parts is not a date."""
        self.assertIsNone(ft._try_parse_date("2026-07-01-extra"))


# ---------------------------------------------------------------------------
# _find_matches — longest-match-wins
# ---------------------------------------------------------------------------

class TestFindMatches(unittest.TestCase):
    """_find_matches: left-to-right, non-overlapping, longest wins."""

    def test_simple_match(self):
        matches = ft._find_matches("hello world", ["world"])
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0], (6, 11, "world"))

    def test_no_match(self):
        matches = ft._find_matches("hello", ["xyz"])
        self.assertEqual(matches, [])

    def test_longest_wins(self):
        """{{A}} should win over stray {{ — longest match at each position."""
        matches = ft._find_matches("hello {{A}} world", ["{{A}}", "{{"])
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0][2], "{{A}}")

    def test_multiple_matches(self):
        matches = ft._find_matches("aaa bbb aaa", ["aaa"])
        self.assertEqual(len(matches), 2)
        self.assertEqual(matches[0][0], 0)
        self.assertEqual(matches[1][0], 8)

    def test_empty_find_ignored(self):
        """Empty find strings are skipped (falsy)."""
        matches = ft._find_matches("hello", ["", "hello"])
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0][2], "hello")


# ---------------------------------------------------------------------------
# tokenise — .docx path
# ---------------------------------------------------------------------------

class TestTokeniseDocx(unittest.TestCase):
    """tokenise: hit counts, not_found, formatting preservation."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.master = Path(self.tmp) / "master.docx"

    def test_basic_tokenise(self):
        _make_docx(self.master, [
            "Dear Ms Jordan Lee,",
            "We confirm your subscription of $1,000,000, effective 01 Jul 2026. Your reference is REF-0001.",
        ])
        dst = Path(self.tmp) / "tokenised.docx"
        rep = ft.tokenise(str(self.master), str(dst), [
            {"find": "Ms Jordan Lee", "token": "Name"},
            {"find": "$1,000,000", "token": "Amount"},
            {"find": "01 Jul 2026", "token": "EffectiveDate"},
            {"find": "REF-0001", "token": "Reference"},
        ])
        self.assertEqual(rep["not_found"], [])
        self.assertEqual(rep["hits"]["Name"], 1)
        self.assertEqual(rep["hits"]["Amount"], 1)
        self.assertEqual(rep["hits"]["EffectiveDate"], 1)
        self.assertEqual(rep["hits"]["Reference"], 1)

    def test_not_found_reported(self):
        """A phrase that doesn't exist is flagged in not_found."""
        _make_docx(self.master, ["Dear Ms Jordan Lee,"])
        dst = Path(self.tmp) / "tokenised.docx"
        rep = ft.tokenise(str(self.master), str(dst), [
            {"find": "Ms Jordan Lee", "token": "Name"},
            {"find": "NONEXISTENT", "token": "Ghost"},
        ])
        self.assertIn("Ghost", rep["not_found"])
        self.assertEqual(rep["hits"]["Name"], 1)
        self.assertEqual(rep["hits"]["Ghost"], 0)

    def test_repeated_phrase_multiple_hits(self):
        """A phrase that appears multiple times gets counted per occurrence."""
        _make_docx(self.master, [
            "Dear Ms Jordan Lee,",
            "Ms Jordan Lee, your subscription is confirmed.",
        ])
        dst = Path(self.tmp) / "tokenised.docx"
        rep = ft.tokenise(str(self.master), str(dst), [
            {"find": "Ms Jordan Lee", "token": "Name"},
        ])
        self.assertEqual(rep["hits"]["Name"], 2)
        self.assertEqual(rep["not_found"], [])

    def test_header_footer_tokenised(self):
        """Tokens in header/footer are found and replaced."""
        _make_docx(self.master, ["Body text."], header="Header: REF-0001", footer="Footer: REF-0001")
        dst = Path(self.tmp) / "tokenised.docx"
        rep = ft.tokenise(str(self.master), str(dst), [
            {"find": "REF-0001", "token": "Reference"},
        ])
        self.assertEqual(rep["hits"]["Reference"], 2)  # header + footer

    def test_bold_preserved(self):
        """The tokenised position keeps bold formatting."""
        _make_docx_with_bold(self.master, "Dear ", "Ms Jordan Lee", ",")
        dst = Path(self.tmp) / "tokenised.docx"
        rep = ft.tokenise(str(self.master), str(dst), [
            {"find": "Ms Jordan Lee", "token": "Name"},
        ])
        self.assertEqual(rep["hits"]["Name"], 1)
        # Read back and check the tokenised run is bold
        doc = Document(str(dst))
        p = doc.paragraphs[0]
        bold_runs = [r for r in p.runs if r.text == "{{Name}}"]
        self.assertEqual(len(bold_runs), 1)
        self.assertTrue(bold_runs[0].bold)

    def test_longest_find_first(self):
        """Overlapping finds: longer phrase wins, shorter reports not_found."""
        _make_docx(self.master, ["Dear Ms Jordan Lee,"])
        dst = Path(self.tmp) / "tokenised.docx"
        rep = ft.tokenise(str(self.master), str(dst), [
            {"find": "Ms Jordan", "token": "ShortName"},
            {"find": "Ms Jordan Lee", "token": "FullName"},
        ])
        self.assertEqual(rep["hits"]["FullName"], 1)
        self.assertEqual(rep["hits"]["ShortName"], 0)
        self.assertIn("ShortName", rep["not_found"])
        # Verify the text
        tokens = ft.tokens_in(str(dst))
        self.assertIn("FullName", tokens)


# ---------------------------------------------------------------------------
# tokenise — .xlsx path
# ---------------------------------------------------------------------------

class TestTokeniseXlsx(unittest.TestCase):
    """tokenise: .xlsx tokenisation + formula preservation."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.master = Path(self.tmp) / "form.xlsx"

    def test_basic_xlsx_tokenise(self):
        _make_xlsx(self.master, {"B5": "{{Name}}", "B6": "{{Amount}}"})
        dst = Path(self.tmp) / "tokenised.xlsx"
        rep = ft.tokenise(str(self.master), str(dst), [
            {"find": "{{Name}}", "token": "Name"},
            {"find": "{{Amount}}", "token": "Amount"},
        ])
        self.assertEqual(rep["not_found"], [])
        self.assertEqual(rep["hits"]["Name"], 1)
        self.assertEqual(rep["hits"]["Amount"], 1)


# ---------------------------------------------------------------------------
# tokens_in
# ---------------------------------------------------------------------------

class TestTokensIn(unittest.TestCase):
    """tokens_in: lists distinct tokens in first-seen order."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_distinct_tokens(self):
        master = Path(self.tmp) / "master.docx"
        _make_docx(master, ["Dear {{Name}}, amount {{Amount}}, ref {{Reference}}"])
        tokens = ft.tokens_in(str(master))
        self.assertEqual(tokens, ["Name", "Amount", "Reference"])

    def test_repeated_token_once(self):
        master = Path(self.tmp) / "master.docx"
        _make_docx(master, ["{{Name}} says {{Name}} is here"])
        tokens = ft.tokens_in(str(master))
        self.assertEqual(tokens, ["Name"])

    def test_whitespace_in_braces(self):
        master = Path(self.tmp) / "master.docx"
        _make_docx(master, ["Dear {{ Name }}, amount {{Amount}}"])
        tokens = ft.tokens_in(str(master))
        self.assertEqual(set(tokens), {"Name", "Amount"})


# ---------------------------------------------------------------------------
# load_rows — data loading
# ---------------------------------------------------------------------------

class TestLoadRows(unittest.TestCase):
    """load_rows: CSV and XLSX data, blank-row skipping."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_csv(self):
        path = Path(self.tmp) / "data.csv"
        _make_data_csv(path, ["Name", "Amount"], [["Alice", "100"], ["Bob", "200"]])
        headers, rows = ft.load_rows(str(path))
        self.assertEqual(headers, ["Name", "Amount"])
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]["Name"], "Alice")
        self.assertEqual(rows[1]["Amount"], "200")

    def test_csv_skips_blank_rows(self):
        path = Path(self.tmp) / "data.csv"
        _make_data_csv(path, ["Name", "Amount"], [
            ["Alice", "100"],
            ["", ""],          # blank — skip
            ["Bob", "200"],
        ])
        _, rows = ft.load_rows(str(path))
        self.assertEqual(len(rows), 2)

    def test_xlsx(self):
        path = Path(self.tmp) / "data.xlsx"
        _make_data_xlsx(path, ["Name", "Amount"], [["Alice", 100], ["Bob", 200]])
        headers, rows = ft.load_rows(str(path))
        self.assertEqual(headers, ["Name", "Amount"])
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]["Name"], "Alice")

    def test_xlsx_datetime_preserved(self):
        """xlsx datetime cells load as datetime objects (so _fmt_value can format them)."""
        path = Path(self.tmp) / "data.xlsx"
        _make_data_xlsx(path, ["Date"], [[datetime.date(2026, 7, 1)]])
        _, rows = ft.load_rows(str(path))
        self.assertIsInstance(rows[0]["Date"], (datetime.datetime, datetime.date))


# ---------------------------------------------------------------------------
# generate — the full pipeline
# ---------------------------------------------------------------------------

class TestGenerate(unittest.TestCase):
    """generate: file production, MISSING flag, unmapped tokens, filename collision."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        # Build a tokenised .docx template
        self.tmpl = Path(self.tmp) / "letter_tokenised.docx"
        _make_docx(self.tmpl, [
            "Dear {{Name}},",
            "We confirm your subscription of {{Amount}}, effective {{EffectiveDate}}. Your reference is {{Reference}}.",
        ])

    def test_basic_generate(self):
        data = Path(self.tmp) / "data.csv"
        _make_data_csv(data, ["Recipient", "Amount", "EffectiveDate", "Reference"],
                       [["Ms Jordan Lee", "1000000", "2026-07-01", "REF-0001"]])
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             token_to_column={"Name": "Recipient"},
                             outdir=str(Path(self.tmp) / "out"),
                             name_pattern="Letter_{Recipient}")
        self.assertEqual(len(report["written"]), 1)
        self.assertEqual(report["unmapped_tokens"], [])
        self.assertEqual(report["written"][0]["missing"], [])

    def test_missing_flag(self):
        """A blank cell produces a visible MISSING flag."""
        data = Path(self.tmp) / "data.csv"
        _make_data_csv(data, ["Recipient", "Amount", "EffectiveDate", "Reference"],
                       [["Ms Jordan Lee", "", "2026-07-01", "REF-0001"]])  # Amount blank
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             token_to_column={"Name": "Recipient"},
                             outdir=str(Path(self.tmp) / "out_missing"),
                             name_pattern="Letter_{Recipient}")
        self.assertIn("Amount", report["written"][0]["missing"])
        # Verify the file contains the MISSING flag
        doc = Document(report["written"][0]["file"])
        text = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("«MISSING: Amount»", text)

    def test_unmapped_token(self):
        """A token with no matching column is flagged as unmapped and MISSING."""
        data = Path(self.tmp) / "data.csv"
        # Data has no Amount column
        _make_data_csv(data, ["Recipient", "EffectiveDate", "Reference"],
                       [["Ms Jordan Lee", "2026-07-01", "REF-0001"]])
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             token_to_column={"Name": "Recipient"},
                             outdir=str(Path(self.tmp) / "out_unmapped"),
                             name_pattern="Letter_{Recipient}")
        self.assertIn("Amount", report["unmapped_tokens"])
        self.assertIn("Amount", report["written"][0]["missing"])

    def test_zero_value_not_missing(self):
        """Amount=0 is a real value — shows as 0, not MISSING."""
        data = Path(self.tmp) / "data.csv"
        _make_data_csv(data, ["Recipient", "Amount", "EffectiveDate", "Reference"],
                       [["Test Person", "0", "2026-07-01", "REF-0001"]])
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             token_to_column={"Name": "Recipient"},
                             outdir=str(Path(self.tmp) / "out_zero"),
                             name_pattern="Letter_{Recipient}")
        self.assertNotIn("Amount", report["written"][0]["missing"])
        doc = Document(report["written"][0]["file"])
        text = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("subscription of 0", text)

    def test_csv_date_formatted(self):
        """CSV date strings render as DD MMM YYYY (the #1 fix)."""
        data = Path(self.tmp) / "data.csv"
        _make_data_csv(data, ["Recipient", "Amount", "EffectiveDate", "Reference"],
                       [["Test Person", "1000000", "2026-07-01", "REF-0001"]])
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             token_to_column={"Name": "Recipient"},
                             outdir=str(Path(self.tmp) / "out_date"),
                             name_pattern="Letter_{Recipient}")
        doc = Document(report["written"][0]["file"])
        text = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("01 Jul 2026", text)

    def test_filename_collision(self):
        """Two rows producing the same filename get deduplicated with (2), (3)."""
        data = Path(self.tmp) / "data.csv"
        _make_data_csv(data, ["Recipient", "Amount", "EffectiveDate", "Reference"],
                       [["Same Name", "100", "2026-07-01", "REF-0001"],
                        ["Same Name", "200", "2026-07-01", "REF-0002"]])
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             token_to_column={"Name": "Recipient"},
                             outdir=str(Path(self.tmp) / "out_collide"),
                             name_pattern="Letter_{Recipient}")
        # Both written, neither skipped
        self.assertEqual(len(report["written"]), 2)
        self.assertEqual(report["skipped"], [])
        # Check filenames are unique
        names = [Path(w["file"]).stem for w in report["written"]]
        self.assertEqual(len(set(names)), 2)

    def test_name_pattern_missing_column_skips(self):
        """If name_pattern references a column not in data, row is skipped."""
        data = Path(self.tmp) / "data.csv"
        _make_data_csv(data, ["Recipient", "Amount", "EffectiveDate", "Reference"],
                       [["Alice", "100", "2026-07-01", "REF-0001"]])
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             token_to_column={"Name": "Recipient"},
                             outdir=str(Path(self.tmp) / "out_skip"),
                             name_pattern="Letter_{NonExistent}")
        self.assertEqual(len(report["skipped"]), 1)
        self.assertEqual(len(report["written"]), 0)

    def test_default_name_pattern(self):
        """Without name_pattern, files are named <template-stem>_<n>."""
        data = Path(self.tmp) / "data.csv"
        _make_data_csv(data, ["Recipient", "Amount", "EffectiveDate", "Reference"],
                       [["Alice", "100", "2026-07-01", "REF-0001"]])
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             token_to_column={"Name": "Recipient"},
                             outdir=str(Path(self.tmp) / "out_default"))
        # Template stem is "letter_tokenised" → "letter" (strips _tokenised)
        fname = Path(report["written"][0]["file"]).name
        self.assertTrue(fname.startswith("letter_") or fname.startswith("Letter_"))


# ---------------------------------------------------------------------------
# generate — .xlsx template path
# ---------------------------------------------------------------------------

class TestGenerateXlsx(unittest.TestCase):
    """generate: .xlsx template — formula preservation."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.tmpl = Path(self.tmp) / "form_tokenised.xlsx"
        _make_xlsx(self.tmpl, {"B5": "{{Name}}", "B6": "{{Amount}}"},
                   formula_cell=("B10", "=B6*1.1"))

    def test_xlsx_generate_preserves_formula(self):
        data = Path(self.tmp) / "data.csv"
        _make_data_csv(data, ["Name", "Amount"], [["Alice", "1000000"]])
        _, rows = ft.load_rows(str(data))
        report = ft.generate(str(self.tmpl), rows,
                             outdir=str(Path(self.tmp) / "out_xlsx"),
                             name_pattern="Form_{Name}")
        wb = openpyxl.load_workbook(report["written"][0]["file"])
        ws = wb.active
        self.assertEqual(ws["B5"].value, "Alice")
        self.assertEqual(ws["B6"].value, "1000000")
        # Formula preserved
        self.assertEqual(ws["B10"].value, "=B6*1.1")


# ---------------------------------------------------------------------------
# _safe_name — filename sanitisation
# ---------------------------------------------------------------------------

class TestSafeName(unittest.TestCase):
    """_safe_name: strips illegal filesystem characters."""

    def test_strips_illegal_chars(self):
        self.assertEqual(ft._safe_name('hello:world/test'), "hello_world_test")

    def test_collapses_whitespace(self):
        self.assertEqual(ft._safe_name("hello   world"), "hello world")

    def test_empty_returns_output(self):
        self.assertEqual(ft._safe_name(""), "output")

    def test_strips_dots(self):
        self.assertEqual(ft._safe_name("..test.."), "test")


# ---------------------------------------------------------------------------
# _resolve_map — token→column mapping
# ---------------------------------------------------------------------------

class TestResolveMap(unittest.TestCase):
    """_resolve_map: defaults to token==header (case-insensitive), overrides."""

    def test_case_insensitive_default(self):
        mapping = ft._resolve_map(["Name", "Amount"], None, ["name", "amount", "extra"])
        self.assertEqual(mapping["Name"], "name")
        self.assertEqual(mapping["Amount"], "amount")

    def test_explicit_override(self):
        mapping = ft._resolve_map(["Name"], {"Name": "Recipient"}, ["Recipient", "Name"])
        self.assertEqual(mapping["Name"], "Recipient")

    def test_no_match(self):
        """Token with no matching column is not in the mapping (→ unmapped)."""
        mapping = ft._resolve_map(["Ghost"], None, ["Name", "Amount"])
        self.assertNotIn("Ghost", mapping)


if __name__ == "__main__":
    unittest.main(verbosity=2)